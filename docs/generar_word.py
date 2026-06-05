"""
Generador Word (.docx) profesional - SoftGold
Genera el Informe de Pruebas de Rendimiento.
Requiere: pip install python-docx
"""

import os, re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Colores como tuplas (r, g, b) ─────────────────────────────────────────
C_PRIMARY     = (0,   100,  60)
C_SECONDARY   = (30,   55,  90)
C_DARK_BG     = (20,   28,  42)
C_TBL_HEADER  = (30,   55,  90)
C_TBL_ZEBRA   = (240, 247, 252)
C_CODE_BG     = (245, 245, 245)
C_WARN_BG     = (255, 248, 220)
C_NOTE_BG     = (232, 244, 255)
C_WHITE       = (255, 255, 255)
C_TEXT        = (30,   30,  30)
C_GRAY        = (110, 110, 110)
C_COVER_SUB   = (180, 220, 200)
C_COVER_DESC  = (160, 200, 190)
C_IMG_BG      = (230, 240, 255)
C_IMG_TEXT    = (80,   80, 150)


def rgb(t):
    """Tupla (r,g,b) → RGBColor."""
    return RGBColor(t[0], t[1], t[2])


def hex_color(t) -> str:
    """Tupla (r,g,b) → string 'RRGGBB' para XML."""
    return f"{t[0]:02X}{t[1]:02X}{t[2]:02X}"


# ── XML helpers ───────────────────────────────────────────────────────────

def cell_bg(cell, color_tuple):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color(color_tuple))
    tcPr.append(shd)


def para_bg(p, color_tuple):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color(color_tuple))
    pPr.append(shd)


def para_border(p, side="left", color="00643C", sz=18, space=6):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"),   "single")
    el.set(qn("w:sz"),    str(sz))
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)


def add_page_number(paragraph):
    run = paragraph.add_run()
    for ftype, text in [("begin", None), (None, "PAGE"), ("end", None)]:
        if ftype:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), ftype)
            run._r.append(el)
        else:
            el = OxmlElement("w:instrText")
            el.text = text
            run._r.append(el)


def add_run(para, text, bold=False, italic=False, size=10.5,
            color=C_TEXT, font="Calibri"):
    r = para.add_run(str(text))
    r.font.name      = font
    r.font.size      = Pt(size)
    r.font.bold      = bold
    r.font.italic    = italic
    r.font.color.rgb = rgb(color)
    return r


# ── Estilos ───────────────────────────────────────────────────────────────

def setup_document(doc: Document, doc_title: str):
    """Configura márgenes, estilos y encabezado/pie."""
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)
        sec.different_first_page_header_footer = True

    # Normal
    n = doc.styles["Normal"]
    n.font.name       = "Calibri"
    n.font.size       = Pt(10.5)
    n.font.color.rgb  = rgb(C_TEXT)
    n.paragraph_format.space_after  = Pt(5)
    n.paragraph_format.space_before = Pt(2)

    # Heading 1
    h = doc.styles["Heading 1"]
    h.font.name       = "Calibri"
    h.font.size       = Pt(13)
    h.font.bold       = True
    h.font.color.rgb  = rgb(C_WHITE)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(8)

    # Heading 2
    h = doc.styles["Heading 2"]
    h.font.name       = "Calibri"
    h.font.size       = Pt(12)
    h.font.bold       = True
    h.font.color.rgb  = rgb(C_SECONDARY)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(5)

    # Heading 3
    h = doc.styles["Heading 3"]
    h.font.name       = "Calibri"
    h.font.size       = Pt(11)
    h.font.bold       = True
    h.font.color.rgb  = rgb(C_PRIMARY)
    h.paragraph_format.space_before = Pt(9)
    h.paragraph_format.space_after  = Pt(4)

    # Encabezado (páginas 2+)
    header = doc.sections[0].header
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_border(hp, side="bottom", color="00643C", sz=10, space=2)
    add_run(hp, f"SoftGold  |  {doc_title}", bold=True,
            size=8, color=C_SECONDARY)

    # Pie de página (páginas 2+)
    footer = doc.sections[0].footer
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_border(fp, side="top", color="00643C", sz=6, space=2)
    add_run(fp, "Uso interno confidencial  |  Pagina ", size=8, color=C_GRAY)
    add_page_number(fp)
    add_run(fp, "  |  softgold.empresa.com", size=8, color=C_GRAY)


# ── Portada ───────────────────────────────────────────────────────────────

def add_cover(doc: Document, subtitle: str, meta_lines: list):
    # Bloque oscuro superior (tabla 1×1)
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.cell(0, 0)
    c.width = Cm(17)
    cell_bg(c, C_DARK_BG)

    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after  = Pt(4)
    add_run(p, "SOFTGOLD", bold=True, size=34, color=C_WHITE)

    p2 = c.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after  = Pt(4)
    add_run(p2, subtitle, bold=True, size=16, color=C_COVER_SUB)

    p3 = c.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after  = Pt(26)
    add_run(p3, "Sistema de Gestion Integral Minera",
            italic=True, size=10.5, color=C_COVER_DESC)

    doc.add_paragraph()

    # Metadatos
    active = [l for l in meta_lines if l]
    if active:
        mt = doc.add_table(rows=len(active), cols=1)
        for idx, line in enumerate(active):
            cell_m = mt.rows[idx].cells[0]
            cell_bg(cell_m, C_NOTE_BG)
            pm = cell_m.paragraphs[0]
            pm.paragraph_format.space_before = Pt(3)
            pm.paragraph_format.space_after  = Pt(3)
            pm.paragraph_format.left_indent  = Cm(0.5)
            if ":" in line:
                key, _, val = line.partition(":")
                add_run(pm, key + ":", bold=True, size=10, color=C_SECONDARY)
                add_run(pm, val,       bold=False, size=10, color=C_TEXT)
            else:
                add_run(pm, line, size=10, color=C_TEXT)

    doc.add_paragraph()

    # Pie de portada
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_bg(pf, C_PRIMARY)
    pf.paragraph_format.space_before = Pt(4)
    pf.paragraph_format.space_after  = Pt(4)
    add_run(pf,
            "Universidad Autonoma de Manizales  |  Ingenieria de Software II  |  2026",
            italic=True, size=9, color=C_WHITE)

    doc.add_page_break()


# ── Elementos de contenido ────────────────────────────────────────────────

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    para_bg(p, C_SECONDARY)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    para_border(p, side="bottom", color="1E375A", sz=8, space=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_body(doc, text):
    p = doc.add_paragraph(style="Normal")
    _inline(p, text)


def add_bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    _inline(p, text)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    _inline(p, text)


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        para_bg(p, C_CODE_BG)
        para_border(p, side="left", color="1E375A", sz=18, space=6)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.6)
        add_run(p, line if line else " ",
                font="Courier New", size=8.5, color=(40, 40, 120))


def add_note(doc, text, kind="NOTE"):
    bg     = C_WARN_BG if kind == "WARN" else C_NOTE_BG
    label  = "ADVERTENCIA" if kind == "WARN" else "NOTA"
    bcolor = "E6A000"   if kind == "WARN" else "00643C"
    lcol   = (180, 100, 0) if kind == "WARN" else C_PRIMARY
    p = doc.add_paragraph()
    para_bg(p, bg)
    para_border(p, side="left", color=bcolor, sz=20, space=6)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.4)
    add_run(p, f"[{label}]  ", bold=True, size=10, color=lcol)
    add_run(p, text.strip(), size=10, color=C_TEXT)


def add_img_placeholder(doc, label):
    p = doc.add_paragraph()
    para_bg(p, C_IMG_BG)
    para_border(p, side="left", color="9090CC", sz=12, space=6)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    add_run(p, f"[ Imagen: {label} ]",
            italic=True, size=10, color=C_IMG_TEXT)


def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1E375A")
    pBdr.append(bot)
    pPr.append(pBdr)


def add_table(doc, rows_data):
    if not rows_data:
        return
    max_cols = max(len(r) for r in rows_data)
    for r in rows_data:
        while len(r) < max_cols:
            r.append("")

    tbl = doc.add_table(rows=len(rows_data), cols=max_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_w = Cm(16.5 / max_cols)

    for ri, row in enumerate(rows_data):
        is_hdr  = (ri == 0)
        is_alt  = (ri % 2 == 0) and not is_hdr
        bg      = C_TBL_HEADER if is_hdr else (C_TBL_ZEBRA if is_alt else C_WHITE)
        txt_col = C_WHITE if is_hdr else C_TEXT

        for ci, cell_txt in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.width = col_w
            cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            add_run(p, str(cell_txt).strip(),
                    bold=is_hdr, size=9, color=txt_col)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Formato inline ────────────────────────────────────────────────────────

def _inline(para, text: str):
    """Añade texto con negrita, cursiva y código inline."""
    parts = re.split(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("***") and part.endswith("***"):
            add_run(para, part[3:-3], bold=True,  italic=True)
        elif part.startswith("**") and part.endswith("**"):
            add_run(para, part[2:-2], bold=True)
        elif part.startswith("*") and part.endswith("*"):
            add_run(para, part[1:-1], italic=True)
        elif part.startswith("`") and part.endswith("`"):
            add_run(para, part[1:-1], font="Courier New",
                    size=9.5, color=(50, 50, 150))
        else:
            add_run(para, part)


# ── Limpieza de texto Markdown ────────────────────────────────────────────

def clean(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)  # links
    text = text.replace("---", "")
    subs = {
        "→":"->","←":"<-","✔":"[OK]","✘":"[X]","■":"-","●":"*",
        "≥":">=","≤":"<=","≠":"!=","×":"x","…":"...",
        "​":"",
        "’":"'","‘":"'","“":'"',"”":'"',
        "–":"-","—":"--",
        "™":"TM","®":"(R)","©":"(C)",
        "☰":"[menu]","▶":">","◦":"o",
        "✏️":"[editar]","🗑️":"[eliminar]",
        "☐":"[ ]","☑":"[x]",
    }
    for ch, r in subs.items():
        text = text.replace(ch, r)
    return text.strip()


# ── Parser Markdown → DOCX ────────────────────────────────────────────────

def parse_md(doc: Document, text: str):
    lines = text.split("\n")
    i = 0
    code_buf  = []
    in_code   = False
    table_buf = []
    in_table  = False

    while i < len(lines):
        raw      = lines[i]
        stripped = raw.strip()

        # Código
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_buf); code_buf = []; in_code = False
            else:
                if in_table:
                    flush_table(doc, table_buf); table_buf = []; in_table = False
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(raw); i += 1; continue

        # Tabla
        if "|" in stripped and stripped.startswith("|"):
            if not in_table: in_table = True; table_buf = []
            table_buf.append(stripped); i += 1; continue
        else:
            if in_table:
                flush_table(doc, table_buf); table_buf = []; in_table = False

        # HR
        if re.match(r"^-{3,}$", stripped):
            add_hr(doc); i += 1; continue

        # Headings
        if stripped.startswith("### "):
            add_h3(doc, clean(stripped[4:])); i += 1; continue
        if stripped.startswith("## "):
            add_h2(doc, clean(stripped[3:])); i += 1; continue
        if stripped.startswith("# "):
            add_h1(doc, clean(stripped[2:])); i += 1; continue

        # Notas
        if stripped.startswith(">"):
            inner = re.sub(r"^>\s*\*\*[^*]+\*\*:?\s*", "", stripped).strip()
            inner = stripped.lstrip("> ").strip() if not inner else inner
            kind  = "WARN" if any(w in stripped for w in
                                  ["ADVERTENCIA","IMPORTANTE"]) else "NOTE"
            add_note(doc, clean(inner), kind); i += 1; continue

        # Capturas
        if "CAPTURA" in stripped or "INSERTAR CAPTURA" in stripped:
            label = re.sub(r"[-\[\]*#]", "", stripped).strip()
            add_img_placeholder(doc, label); i += 1; continue

        # Lista numerada
        nm = re.match(r"^\d+\.\s+(.*)", stripped)
        if nm:
            add_numbered(doc, clean(nm.group(1))); i += 1; continue

        # Lista viñeta
        bm = re.match(r"^[-*]\s+(.*)", stripped)
        if bm:
            lvl = 1 if raw.startswith("  ") else 0
            add_bullet(doc, clean(bm.group(1)), lvl); i += 1; continue

        # Párrafo
        if stripped:
            add_body(doc, clean(stripped)); i += 1; continue

        i += 1  # línea en blanco

    if in_code:   add_code_block(doc, code_buf)
    if in_table:  flush_table(doc, table_buf)


def flush_table(doc, rows):
    data = [r for r in rows if not re.match(r"^\|[\s\-|:]+\|?$", r)]
    if not data: return
    parsed = []
    for row in data:
        cells = [clean(c.strip()) for c in row.strip("|").split("|")]
        parsed.append(cells)
    if parsed:
        mc = max(len(r) for r in parsed)
        for r in parsed:
            while len(r) < mc: r.append("")
        add_table(doc, parsed)


# ── Función principal ─────────────────────────────────────────────────────

def generate(md_file, docx_file, title, subtitle, meta_lines):
    md_path   = os.path.join(DOCS_DIR, md_file)
    docx_path = os.path.join(DOCS_DIR, docx_file)

    print(f"\n{'='*55}")
    print(f"  Generando: {docx_file}")
    print(f"{'='*55}")

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    setup_document(doc, title)
    add_cover(doc, subtitle, meta_lines)
    parse_md(doc, content)
    doc.save(docx_path)

    size_kb = os.path.getsize(docx_path) / 1024
    print(f"  [OK] {docx_path}  ({size_kb:.1f} KB)")


# ── Punto de entrada ──────────────────────────────────────────────────────

if __name__ == "__main__":
    print("\n" + "="*55)
    print("  GENERADOR WORD - SOFTGOLD")
    print("="*55)

    generate(
        md_file    = "informe_pruebas_rendimiento.md",
        docx_file  = "SoftGold_Informe_Pruebas_Rendimiento.docx",
        title      = "INFORME DE PRUEBAS DE RENDIMIENTO",
        subtitle   = "Informe de Pruebas de Rendimiento, Carga y Estres",
        meta_lines = [
            "Version: 1.0.0",
            "Fecha: Mayo 2026",
            "Dirigido a: Equipo QA | DevOps | Arquitectura",
            "Estado: Documento final aprobado",
            "",
            "Pruebas: Carga | Estres | Estabilidad | Concurrencia",
            "Herramientas: JMeter 5.6 | JVisualVM | MySQL EXPLAIN",
        ]
    )

    generate(
        md_file    = "manual_tecnico.md",
        docx_file  = "SoftGold_Manual_Tecnico.docx",
        title      = "MANUAL TECNICO DEL SISTEMA",
        subtitle   = "Manual Tecnico del Sistema",
        meta_lines = [
            "Version: 1.0.0",
            "Fecha: Mayo 2026",
            "Dirigido a: Desarrolladores y equipo tecnico",
            "Estado: Documento oficial",
            "",
            "Tecnologias: Spring Boot 3.4.4 | Java 17 | MySQL 8 | Thymeleaf",
            "Puerto: 9090  |  BD: softgold  |  Puerto MySQL: 3306",
        ]
    )

    generate(
        md_file    = "manual_usuario.md",
        docx_file  = "SoftGold_Manual_Usuario.docx",
        title      = "MANUAL DE USUARIO",
        subtitle   = "Manual de Usuario",
        meta_lines = [
            "Version: 1.0.0",
            "Fecha: Mayo 2026",
            "Dirigido a: Usuarios finales del sistema",
            "Estado: Documento oficial",
            "",
            "Modulos: Minas | Mapas | Riesgos | Exploracion",
            "         Foro | Informes | Soporte | Mi Cuenta",
        ]
    )

    generate(
        md_file    = "manual_implementacion.md",
        docx_file  = "SoftGold_Manual_Implementacion.docx",
        title      = "MANUAL DE IMPLEMENTACION Y DESPLIEGUE",
        subtitle   = "Manual de Implementacion y Despliegue",
        meta_lines = [
            "Version: 1.0.0",
            "Fecha: Mayo 2026",
            "Dirigido a: Equipo DevOps / Administradores",
            "Estado: Documento oficial",
            "",
            "Ambientes: Desarrollo | QA | Produccion",
            "Infraestructura: Nginx | MySQL | systemd | Docker",
        ]
    )

    print("\n" + "="*55)
    print("  COMPLETADO")
    print(f"  Archivos en: {DOCS_DIR}")
    print("="*55 + "\n")
